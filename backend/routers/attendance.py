from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
import io
from backend.schemas import users as user_schemas
from backend.crud import attendance as attendance_crud
from backend.crud import classes as class_crud
from backend.core import database, security

from backend.schemas import attendance as attendance_schemas

router = APIRouter()

@router.put("/classes/{class_id}/attendance/{session_id}", response_model=attendance_schemas.AttendanceSession)
def update_attendance_session(
    class_id: int, 
    session_id: int, 
    session: attendance_schemas.AttendanceSessionCreate, 
    db: Session = Depends(database.get_db), 
    current_user: user_schemas.User = Depends(security.get_current_user)
):
    # Verify ownership
    db_class = class_crud.get_class(db, class_id=class_id)
    if not db_class or db_class.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    updated_session = attendance_crud.update_attendance_session(db, session_id=session_id, session_data=session)
    if not updated_session:
        raise HTTPException(status_code=404, detail="Session not found")
    return updated_session

@router.delete("/classes/{class_id}/attendance/{session_id}")
def delete_attendance_session(
    class_id: int, 
    session_id: int, 
    db: Session = Depends(database.get_db), 
    current_user: user_schemas.User = Depends(security.get_current_user)
):
    # Verify ownership
    db_class = class_crud.get_class(db, class_id=class_id)
    if not db_class or db_class.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    deleted_session = attendance_crud.delete_attendance_session(db, session_id=session_id)
    if not deleted_session:
        raise HTTPException(status_code=404, detail="Session not found")
    return {"message": "Session deleted successfully"}

@router.get("/attendance-sessions/{session_id}")
def read_attendance_session(session_id: int, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    session = attendance_crud.get_attendance_session(db, session_id=session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    # Verify ownership via class... lazy way: fetch class
    db_class = class_crud.get_class(db, class_id=session.class_id)
    if db_class.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    return session

@router.get("/attendance-sessions/{session_id}/report/docx")
def generate_session_report(session_id: int, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    session = attendance_crud.get_attendance_session(db, session_id=session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Verify ownership
    db_class = class_crud.get_class(db, class_id=session.class_id)
    if db_class.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    document = Document()
    document.add_heading(f'Relatório de Aula', 0)
    document.add_paragraph(f'Turma: {db_class.name}')
    document.add_paragraph(f'Data: {session.date}')
    document.add_paragraph(f'Descrição: {session.description}')
    
    document.add_heading('Frequência e Notas', level=1)
    
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aluno'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Nota'
    hdr_cells[3].text = 'Observação'
    
    for log in session.logs:
        row_cells = table.add_row().cells
        # log.student might be lazy loaded, proceed with caution or let ORM handle
        student_name = log.student.name if log.student else "Unknown"
        row_cells[0].text = student_name
        row_cells[1].text = "Presente" if log.status == 'present' else "Ausente"
        row_cells[2].text = str(log.grade) if log.grade is not None else '-'
        row_cells[3].text = str(log.observation) if log.observation else '-'
        
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=aula_{session.date}.docx"}
    )
