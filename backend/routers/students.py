from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
from typing import List, Optional
import io
import pydantic
from backend.schemas import students as student_schemas
from backend.schemas import users as user_schemas
from backend.crud import students as student_crud
from backend.core import database, security

router = APIRouter()

@router.get("/students/", response_model=List[student_schemas.Student])
def read_students(
    skip: int = 0, 
    limit: int = 100, 
    search: Optional[str] = None,
    db: Session = Depends(database.get_db), 
    current_user: user_schemas.User = Depends(security.get_current_user)
):
    return student_crud.get_students(db, user_id=current_user.id, skip=skip, limit=limit, search=search)

@router.post("/students/", response_model=student_schemas.Student)
def create_student(student: student_schemas.StudentCreate, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    return student_crud.create_student(db=db, student=student, user_id=current_user.id)

class StudentUpdate(pydantic.BaseModel):
    name: str

@router.put("/students/{student_id}")
def update_student(student_id: int, student_data: StudentUpdate, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    # Verify ownership logic could be added here (check student -> owner_id)
    student = student_crud.update_student(db, student_id=student_id, name=student_data.name)
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    return student

@router.delete("/students/{student_id}")
def delete_student(student_id: int, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    student = student_crud.delete_student(db, student_id=student_id)
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    return {"detail": "Student deleted"}

@router.get("/students/{student_id}/evolution", response_model=List[student_schemas.StudentEvolutionPoint])
def get_student_evolution(student_id: int, db: Session = Depends(database.get_db), current_user: user_schemas.User = Depends(security.get_current_user)):
    results = student_crud.get_student_evolution(db, student_id=student_id)
    
    response = []
    for log in results:
        response.append({
            "date": log.session.date,
            "grade": log.grade,
            "status": log.status
        })
    return response

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

@router.post("/students/{student_id}/report/docx")
def generate_student_report(
    student_id: int, 
    report_data: student_schemas.StudentReportRequest,
    db: Session = Depends(database.get_db), 
    current_user: user_schemas.User = Depends(security.get_current_user)
):
    stats = student_crud.get_student_report_stats(db, student_id=student_id)
    if not stats:
        raise HTTPException(status_code=404, detail="Student not found")
    
    document = Document()
    
    # Title
    title = document.add_heading(f'Relatório de Desempenho', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student Info Section
    document.add_heading('Informações do Aluno', level=1)
    p = document.add_paragraph()
    p.add_run('Nome: ').bold = True
    p.add_run(f'{stats["student"].name}\n')
    if stats["student"].parent_name:
        p.add_run('Responsável: ').bold = True
        p.add_run(f'{stats["student"].parent_name}\n')
    
    # Statistics Section
    document.add_heading('Resumo de Atividades', level=1)
    table_stats = document.add_table(rows=1, cols=4)
    table_stats.style = 'Table Grid'
    hdr_stats = table_stats.rows[0].cells
    hdr_stats[0].text = 'Total Aulas'
    hdr_stats[1].text = 'Presenças'
    hdr_stats[2].text = 'Frequência'
    hdr_stats[3].text = 'Média Notas'
    
    # Style Headers
    for cell in hdr_stats:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        
    row_stats = table_stats.add_row().cells
    row_stats[0].text = str(stats["total_classes"])
    row_stats[1].text = str(stats["present"])
    row_stats[2].text = f'{stats["attendance_rate"]}%'
    row_stats[3].text = f'{stats["avg_grade"]}'

    # Evolution Chart Section
    if report_data.chart_image:
        document.add_heading('Gráfico de Evolução', level=1)
        try:
            # Decode base64 image
            # Format usually: "data:image/png;base64,....."
            if "," in report_data.chart_image:
                header, encoded = report_data.chart_image.split(",", 1)
            else:
                encoded = report_data.chart_image
            
            image_data = base64.b64decode(encoded)
            image_stream = io.BytesIO(image_data)
            
            # Add picture centered
            document.add_picture(image_stream, width=Inches(6))
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error adding image: {e}")
            document.add_paragraph("[Erro ao incluir o gráfico]")

    
    # Detailed History
    document.add_heading('Histórico Detalhado', level=1)
    
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Conteúdo/Descrição'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Nota'
    hdr_cells[4].text = 'Observação'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    for log in stats["logs"]:
        row_cells = table.add_row().cells
        row_cells[0].text = log.session.date.strftime('%d/%m/%Y')
        row_cells[1].text = str(log.session.description)
        
        # Translate status
        # Translate status
        status_map = {
            'PRESENT': 'Presente', 
            'ABSENT': 'Ausente', 
            'LATE': 'Atrasado', 
            'Justified': 'Justificado',
            'present': 'Presente',
            'absent': 'Ausente',
            'late': 'Atrasado',
            'justified': 'Justificado'
        }
        row_cells[2].text = status_map.get(log.status, log.status)
        
        row_cells[3].text = str(log.grade) if log.grade is not None else '-'
        row_cells[4].text = str(log.observation) if log.observation else '-'
        
    # Save to stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Relatorio_{stats['student'].name.replace(' ', '_')}.docx"}
    )
